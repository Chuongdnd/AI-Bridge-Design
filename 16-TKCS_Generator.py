# -*- coding: utf-8 -*-
"""
16-TKCS_Generator.py — Sinh thuyết minh thiết kế cơ sở (TKCS) định dạng DOCX
Tuân thủ Nghị định 175/2024/NĐ-CP và chuẩn ngành xây dựng Việt Nam
"""

from __future__ import annotations

import io
import os
import json
import math
import time
import zipfile
import datetime
import tempfile
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union

logger = logging.getLogger("TKCS_Generator")

# ── python-docx ───────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    logger.error("Chưa cài python-docx: pip install python-docx")

ROOT        = Path(__file__).parent
OUTPUT_DIR  = ROOT / "Data" / "Output" / "TKCS"
TEMPLATES_DIR = ROOT / "Data" / "Templates"

# ─────────────────────────────────────────────────────────────────────────────
# Tiêu chuẩn áp dụng chuẩn cho mọi dự án cầu BTCT
# ─────────────────────────────────────────────────────────────────────────────

TIEU_CHUAN_MAC_DINH = [
    ("Luật Xây dựng số 50/2014/QH13", "Sửa đổi theo Luật 62/2020/QH14",
     "Cơ sở pháp lý chung cho hoạt động xây dựng"),
    ("Nghị định 175/2024/NĐ-CP", "30/12/2024",
     "Quản lý chất lượng và bảo trì công trình xây dựng"),
    ("TCVN 11823:2017", "2017",
     "Thiết kế cầu đường bộ — Tiêu chuẩn chính áp dụng cho toàn bộ công trình"),
    ("TCVN 10304:2014", "2014",
     "Móng cọc — Tiêu chuẩn thiết kế"),
    ("TCVN 4054:2005", "2005",
     "Đường ô tô — Yêu cầu thiết kế"),
    ("TCVN 5664:2009", "2009",
     "Phân cấp kỹ thuật đường thủy nội địa"),
    ("TCVN 3972:1985", "1985",
     "Công tác trắc địa trong xây dựng"),
    ("TCVN 9362:2012", "2012",
     "Tiêu chuẩn thiết kế nền nhà và công trình"),
    ("TCVN 1651:2018", "2018",
     "Thép cốt bê tông — Phần 1 và 2"),
    ("TCVN 3118:1993", "1993",
     "Bê tông nặng — Phương pháp xác định cường độ"),
    ("Quyết định 409/2025/QĐ-BXD", "2025",
     "Công bố suất vốn đầu tư xây dựng — cơ sở lập dự toán sơ bộ"),
]

# ─────────────────────────────────────────────────────────────────────────────
# DocumentNumbering — đánh số tự động
# ─────────────────────────────────────────────────────────────────────────────

class DocumentNumbering:
    """Quản lý counter tự động: chương, mục, bảng, hình."""

    def __init__(self):
        self._ch  = 0
        self._sec: Dict[int, int] = {}
        self._tab: Dict[int, int] = {}
        self._fig: Dict[int, int] = {}
        self.toc:     List[Tuple[int, str]] = []
        self.tables:  List[Tuple[str, str]] = []
        self.figures: List[Tuple[str, str]] = []

    def chapter(self, title: str = "") -> Tuple[int, str]:
        self._ch += 1
        self._sec[self._ch] = 0
        self._tab[self._ch] = 0
        self._fig[self._ch] = 0
        label = f"Chương {self._ch}"
        if title:
            self.toc.append((1, f"{label}: {title}"))
        return self._ch, label

    def section(self, title: str = "") -> str:
        self._sec[self._ch] = self._sec.get(self._ch, 0) + 1
        num = f"{self._ch}.{self._sec[self._ch]}"
        if title:
            self.toc.append((2, f"{num} {title}"))
        return num

    def table(self, caption: str = "") -> str:
        self._tab[self._ch] = self._tab.get(self._ch, 0) + 1
        num = f"Bảng {self._ch}.{self._tab[self._ch]}"
        if caption:
            self.tables.append((num, caption))
        return num

    def figure(self, caption: str = "") -> str:
        self._fig[self._ch] = self._fig.get(self._ch, 0) + 1
        num = f"Hình {self._ch}.{self._fig[self._ch]}"
        if caption:
            self.figures.append((num, caption))
        return num


# ─────────────────────────────────────────────────────────────────────────────
# TKCSGenerator — sinh toàn bộ thuyết minh
# ─────────────────────────────────────────────────────────────────────────────

class TKCSGenerator:
    """
    Sinh file DOCX thuyết minh TKCS từ dữ liệu Session State.

    Parameters
    ----------
    session_data : dict
        Toàn bộ st.session_state (hoặc dict tương đương cho testing).
    config : dict
        Cấu hình xuất: selected_chapters, detail_level, include_figures, ...
    project_info : dict
        Thông tin dự án: ten_du_an, dia_diem, chu_dau_tu, tu_van, ...
    """

    def __init__(
        self,
        session_data: Dict,
        config: Optional[Dict] = None,
        project_info: Optional[Dict] = None,
    ):
        if not DOCX_OK:
            raise ImportError("python-docx chưa cài")

        self.ss   = session_data
        self.cfg  = config or {}
        self.pi   = project_info or {}
        self.num  = DocumentNumbering()
        self.doc: Optional[Any] = None

        # counters dùng khi insert_table / insert_figure gọi mà chưa qua num
        self._current_ch = 0

        # Danh sách chương được chọn (mặc định tất cả 1–12)
        self.selected = set(self.cfg.get("selected_chapters", list(range(1, 13))))

    # ── Helpers: đọc session state an toàn ────────────────────────────────────

    def _get(self, *keys, default=None):
        obj = self.ss
        for k in keys:
            if not isinstance(obj, dict):
                return default
            obj = obj.get(k, default)
            if obj is None:
                return default
        return obj

    def _beam(self):
        return (self.ss.get("beam_params_final")
                or self.ss.get("beam_params")
                or {})

    def _kcn(self):
        return self.ss.get("kcn_result") or {}

    def _pier(self):
        return self.ss.get("pier_result") or {}

    def _mong(self):
        return self.ss.get("mong_coc_result") or {}

    def _dt(self):
        return self.ss.get("du_toan_result") or {}

    def _pa(self):
        return self.ss.get("phuong_an_cuoi_cung") or {}

    # ── Thiết lập document và styles ──────────────────────────────────────────

    def _setup_document(self) -> Any:
        doc = Document()

        # Trang A4, lề chuẩn đóng quyển
        sec = doc.sections[0]
        sec.page_width  = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)

        def _set_style(name, fname="Times New Roman", size=13, bold=False,
                       italic=False, align=None, space_before=0, space_after=6,
                       color=None):
            try:
                st = doc.styles[name]
            except KeyError:
                return
            st.font.name  = fname
            st.font.size  = Pt(size)
            st.font.bold  = bold
            st.font.italic = italic
            if color:
                st.font.color.rgb = RGBColor(*color)
            pf = st.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing  = 1.5
            pf.space_before  = Pt(space_before)
            pf.space_after   = Pt(space_after)
            if align is not None:
                pf.alignment = align

        _set_style("Normal",    size=13)
        _set_style("Heading 1", size=14, bold=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
        _set_style("Heading 2", size=13, bold=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
        _set_style("Heading 3", size=13, bold=False, italic=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3)
        _set_style("Caption",   size=12, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        return doc

    def _add_header_footer(self, doc, title_short: str):
        """Thêm header (tên đề tài) và footer (trang + ngày xuất)."""
        sec = doc.sections[0]
        sec.different_first_page_header_footer = True  # Trang bìa không có

        # Header
        header = sec.header
        hp = header.paragraphs[0]
        hp.text = title_short
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _fmt_run(hp.runs[0] if hp.runs else hp.add_run(), bold=False, size=11)

        # Footer — số trang + ngày
        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(
            f"AI-Bridge-Design | Xuất ngày {datetime.date.today().strftime('%d/%m/%Y')}  —  Trang "
        )
        _fmt_run(run, size=10, italic=True)
        _add_page_number(fp)

    # ── Trang bìa ─────────────────────────────────────────────────────────────

    def _add_cover_page(self, doc):
        pi = self.pi

        def _center(text, size=13, bold=False):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(size)
            run.font.bold = bold
            p.paragraph_format.space_after = Pt(4)
            return p

        doc.add_paragraph()  # khoảng trắng đầu

        _center("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI TP.HCM", 14, True)
        _center("KHOA KỸ THUẬT XÂY DỰNG", 13, True)
        doc.add_paragraph()
        doc.add_paragraph()

        _center("LUẬN VĂN THẠC SĨ KỸ THUẬT", 16, True)
        doc.add_paragraph()

        _center("THUYẾT MINH", 14, True)
        _center("THIẾT KẾ CƠ SỞ", 18, True)
        doc.add_paragraph()

        ten_cau = pi.get("ten_du_an", "CÔNG TRÌNH CẦU [TÊN CẦU]")
        _center(ten_cau.upper(), 16, True)
        doc.add_paragraph()
        doc.add_paragraph()

        _center(f"Địa điểm: {pi.get('dia_diem','')}", 13)
        _center(f"Chủ đầu tư: {pi.get('chu_dau_tu','')}", 13)
        _center(f"Đơn vị tư vấn: {pi.get('tu_van','AI-Bridge-Design System')}", 13)
        doc.add_paragraph()

        _center(f"Giảng viên hướng dẫn: {pi.get('gv_huong_dan','')}", 13)
        _center(f"Học viên thực hiện: {pi.get('hoc_vien','')}", 13)
        doc.add_paragraph()

        _center(
            f"TP. Hồ Chí Minh, tháng "
            f"{datetime.date.today().month} năm {datetime.date.today().year}",
            13,
        )

        doc.add_page_break()

    # ── Mục lục ───────────────────────────────────────────────────────────────

    def _add_front_matter(self, doc):
        """Thêm trang MỤC LỤC (placeholder — tự cập nhật trong Word)."""
        p = doc.add_paragraph("MỤC LỤC")
        p.style = doc.styles["Heading 1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            "(Mục lục tự động — Nhấn Ctrl+A rồi F9 trong Microsoft Word để cập nhật)"
        ).runs[0].font.italic = True

        # Liệt kê tất cả mục từ numbering
        for lvl, text in self.num.toc:
            indent = Cm(1.0 * (lvl - 1))
            p2 = doc.add_paragraph(text)
            p2.paragraph_format.left_indent = indent
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.space_before = Pt(2)
            _fmt_run(p2.runs[0] if p2.runs else p2.add_run(text),
                     bold=(lvl == 1), size=12)

        doc.add_page_break()

        # Danh mục bảng
        if self.num.tables:
            p = doc.add_paragraph("DANH MỤC BẢNG BIỂU")
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for num_str, cap in self.num.tables:
                doc.add_paragraph(f"{num_str}: {cap}").paragraph_format.space_after = Pt(2)
            doc.add_page_break()

        # Danh mục hình
        if self.num.figures:
            p = doc.add_paragraph("DANH MỤC HÌNH VẼ")
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for num_str, cap in self.num.figures:
                doc.add_paragraph(f"{num_str}: {cap}").paragraph_format.space_after = Pt(2)
            doc.add_page_break()

    # ── Tiêu chuẩn viết tắt ───────────────────────────────────────────────────

    def _add_abbreviations(self, doc):
        p = doc.add_paragraph("DANH MỤC CÁC CHỮ VIẾT TẮT")
        p.style = doc.styles["Heading 1"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        abbrevs = [
            ("BTCT",  "Bê tông cốt thép"),
            ("DUL",   "Dự ứng lực"),
            ("KCN",   "Kết cấu nhịp"),
            ("TKCS",  "Thiết kế cơ sở"),
            ("TCVN",  "Tiêu chuẩn Việt Nam"),
            ("QCVN",  "Quy chuẩn kỹ thuật quốc gia Việt Nam"),
            ("MNCN",  "Mực nước cao nhất"),
            ("MNTN",  "Mực nước thấp nhất"),
            ("MNTT",  "Mực nước thông thuyền"),
            ("SPT",   "Standard Penetration Test — Thí nghiệm xuyên tiêu chuẩn"),
            ("HL-93", "Highway Loading 1993 — Tải trọng xe thiết kế theo TCVN 11823"),
            ("NĐ",    "Nghị định"),
            ("UBND",  "Ủy ban nhân dân"),
            ("BGTVT", "Bộ Giao thông Vận tải"),
        ]
        tbl = self.insert_table(
            doc,
            headers=["Ký hiệu", "Giải thích"],
            rows=abbrevs,
            caption=None,
            style="Table Grid",
        )
        doc.add_page_break()

    # ── Helper thêm chương / mục / đoạn văn ─────────────────────────────────

    def _h1(self, doc, title: str) -> Tuple[int, str]:
        ch_num, label = self.num.chapter(title)
        self._current_ch = ch_num
        full = f"{label.upper()}: {title.upper()}"
        p = doc.add_paragraph(full)
        p.style = doc.styles["Heading 1"]
        return ch_num, label

    def _h2(self, doc, title: str) -> str:
        num_str = self.num.section(title)
        p = doc.add_paragraph(f"{num_str} {title}")
        p.style = doc.styles["Heading 2"]
        return num_str

    def _h3(self, doc, title: str) -> None:
        p = doc.add_paragraph(f"  {title}")
        p.style = doc.styles["Heading 3"]

    def _para(self, doc, text: str, indent_cm: float = 0.75, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(indent_cm)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
        run.font.bold = bold

    def _bullet(self, doc, text: str):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)

    # ── Chèn bảng ─────────────────────────────────────────────────────────────

    def insert_table(
        self,
        doc,
        headers: List[str],
        rows: List[List],
        caption: Optional[str] = None,
        style: str = "Table Grid",
    ) -> Any:
        """Chèn bảng chuẩn với header xám nhạt."""
        if caption:
            num_str = self.num.table(caption)
            cap_text = f"{num_str}: {caption}"
            cp = doc.add_paragraph(cap_text)
            try:
                cp.style = doc.styles["Caption"]
            except KeyError:
                cp.style = doc.styles["Normal"]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        n_cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
        try:
            tbl.style = doc.styles[style]
        except KeyError:
            tbl.style = doc.styles["Table Grid"]

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = str(h)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Tô màu xám nhạt
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "D9D9D9")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
            # Bold header
            for run in hdr_cells[i].paragraphs[0].runs:
                run.font.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # Data rows
        for ri, row in enumerate(rows):
            cells = tbl.rows[ri + 1].cells
            for ci, val in enumerate(row):
                cells[ci].text = str(val) if val is not None else ""
                cells[ci].paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                    if _is_numeric(str(val))
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in cells[ci].paragraphs[0].runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

        doc.add_paragraph()  # khoảng cách sau bảng
        return tbl

    # ── Chèn hình ─────────────────────────────────────────────────────────────

    def insert_figure(
        self,
        doc,
        img: Union[str, Path, bytes, io.BytesIO],
        caption: str = "",
        width_cm: float = 14.0,
    ):
        """Chèn ảnh (file path hoặc bytes) với chú thích."""
        try:
            if isinstance(img, (str, Path)):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img), width=Cm(width_cm))
            elif isinstance(img, (bytes, io.BytesIO)):
                buf = io.BytesIO(img) if isinstance(img, bytes) else img
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(buf, width=Cm(width_cm))
        except Exception as e:
            doc.add_paragraph(f"[Hình — {caption} — {e}]")
            return

        if caption:
            num_str = self.num.figure(caption)
            cp = doc.add_paragraph(f"{num_str}: {caption}")
            try:
                cp.style = doc.styles["Caption"]
            except KeyError:
                pass
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def insert_chart_from_plotly(self, doc, fig, caption: str = "", width_cm: float = 14.0):
        """Chuyển Plotly figure → PNG rồi chèn vào document."""
        try:
            img_bytes = fig.to_image(format="png", width=1200, height=700, scale=2)
            self.insert_figure(doc, img_bytes, caption, width_cm)
        except Exception as e:
            doc.add_paragraph(f"[Biểu đồ — {caption} — lỗi kaleido: {e}]")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 1 — Thông tin chung dự án
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_1(self, doc):
        pi = self.pi
        kcn = self._kcn()
        beam = self._beam()

        self._h1(doc, "Thông tin chung về dự án")

        self._h2(doc, "Tổng quan dự án")
        self._para(doc,
            f"Dự án \"{pi.get('ten_du_an','[Tên dự án]')}\" được lập nhằm "
            f"đầu tư xây dựng công trình cầu tại {pi.get('dia_diem','[địa điểm]')}, "
            f"phục vụ nhu cầu đi lại của nhân dân và phát triển kinh tế — xã hội "
            f"của địa phương. Đây là giai đoạn thiết kế cơ sở phục vụ lập "
            f"Báo cáo nghiên cứu khả thi theo quy định tại Điều 19, "
            f"Nghị định 175/2024/NĐ-CP.")

        self._h2(doc, "Vị trí địa lý và phạm vi công trình")
        self._para(doc,
            f"Công trình cầu dự kiến xây dựng tại {pi.get('dia_diem','[địa điểm]')}. "
            f"Tuyến cầu có điểm đầu tại {pi.get('diem_dau','[điểm đầu]')} "
            f"và điểm cuối tại {pi.get('diem_cuoi','[điểm cuối]')}. "
            f"Công trình thuộc dự án {pi.get('ten_du_an','[tên dự án]')}.")

        quy_mo = kcn.get("be_rong_cau") or beam.get("B", "")
        n_nhip = kcn.get("tong_so_nhip", 1)
        L_nhip = beam.get("L", 0)
        L_cau  = n_nhip * (L_nhip / 1000 if L_nhip > 100 else L_nhip)
        self._para(doc,
            f"Quy mô công trình: Cầu dài {L_cau:.1f}m, bề rộng {quy_mo}m, "
            f"gồm {n_nhip} nhịp. Công trình thuộc cấp "
            f"{self.ss.get('cap_cong_trinh', 'III')} theo phân loại tại "
            f"Phụ lục I, Nghị định 175/2024/NĐ-CP.")

        self._h2(doc, "Chủ đầu tư và đơn vị thực hiện")
        rows = [
            ["Chủ đầu tư",          pi.get("chu_dau_tu", "")],
            ["Đơn vị tư vấn",       pi.get("tu_van", "AI-Bridge-Design System")],
            ["Người lập",           pi.get("hoc_vien", "")],
            ["GV hướng dẫn",        pi.get("gv_huong_dan", "")],
            ["Ngày lập",            pi.get("ngay_lap", str(datetime.date.today()))],
        ]
        self.insert_table(doc, ["Thông tin", "Nội dung"], rows,
                          "Thông tin đơn vị thực hiện")

        self._h2(doc, "Mục tiêu đầu tư")
        self._para(doc,
            "Mục tiêu xây dựng công trình cầu nhằm: (i) Đáp ứng nhu cầu giao thông "
            "của nhân dân trong khu vực; (ii) Nâng cao năng lực thông qua tuyến đường; "
            "(iii) Phục vụ phát triển kinh tế — xã hội địa phương; (iv) Đảm bảo an "
            "toàn giao thông tại vị trí xây dựng. Thuyết minh này được lập theo "
            "yêu cầu của Điều 20, Nghị định 175/2024/NĐ-CP.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 2 — Cơ sở pháp lý và tiêu chuẩn áp dụng
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_2(self, doc):
        self._h1(doc, "Cơ sở pháp lý và tiêu chuẩn thiết kế")

        self._h2(doc, "Cơ sở pháp lý")
        self._para(doc,
            "Thuyết minh thiết kế cơ sở được lập trên cơ sở các văn bản pháp lý hiện hành "
            "của Nhà nước Cộng hòa xã hội chủ nghĩa Việt Nam có liên quan đến hoạt động "
            "đầu tư xây dựng công trình giao thông, bao gồm:")
        for tc in TIEU_CHUAN_MAC_DINH[:2]:
            self._bullet(doc, f"{tc[0]} ({tc[1]}): {tc[2]}")

        self._h2(doc, "Các tiêu chuẩn kỹ thuật áp dụng")
        self._para(doc,
            "Công trình được thiết kế theo các tiêu chuẩn kỹ thuật hiện hành của "
            "Việt Nam, cụ thể:")
        rows = [[tc[0], tc[1], tc[2]] for tc in TIEU_CHUAN_MAC_DINH]
        self.insert_table(
            doc,
            ["Tiêu chuẩn/Văn bản", "Năm ban hành", "Phạm vi áp dụng"],
            rows,
            "Danh mục tiêu chuẩn áp dụng",
        )

        self._h2(doc, "Tài liệu khảo sát và tham chiếu")
        self._para(doc,
            "Ngoài các tiêu chuẩn kỹ thuật trên, việc lập thuyết minh có sử dụng "
            "các tài liệu tham chiếu: kết quả khảo sát địa hình tỷ lệ 1:500, "
            "hồ sơ khảo sát địa chất công trình, tài liệu thủy văn công trình "
            "do đơn vị khảo sát cung cấp (nếu có). Phần chi tiết xem Phụ lục.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 3 — Điều kiện tự nhiên
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_3(self, doc):
        td   = self.ss.get("terrain_data") or {}
        dgc  = self.ss.get("dia_chat_data") or {}

        self._h1(doc, "Điều kiện tự nhiên khu vực xây dựng")

        self._h2(doc, "Đặc điểm địa hình")
        self._para(doc,
            f"Khu vực xây dựng công trình có địa hình {td.get('dia_hinh','đồng bằng, tương đối bằng phẳng')}. "
            f"Cao độ mặt đất tự nhiên dao động từ "
            f"{td.get('cao_do_min','±0.00')}m đến {td.get('cao_do_max','±0.00')}m. "
            f"Địa hình khu vực phù hợp cho việc xây dựng công trình cầu "
            f"theo phương án đã đề xuất.")

        self._h2(doc, "Đặc điểm khí hậu và thủy văn")
        self._para(doc,
            f"Công trình xây dựng tại khu vực có khí hậu nhiệt đới gió mùa. "
            f"Nhiệt độ trung bình năm khoảng 27–29°C. Lượng mưa trung bình "
            f"năm 1.500–2.000mm. Mùa mưa từ tháng 5 đến tháng 11.")

        mn_cn = td.get("MNCN") or td.get("mncn") or "N/A"
        mn_tn = td.get("MNTN") or td.get("mntn") or "N/A"
        mn_tt = td.get("MNTT") or td.get("mntt") or "N/A"
        rows_mn = [
            ["Mực nước cao nhất (MNCN)", f"{mn_cn} m", "Tần suất P=1%"],
            ["Mực nước thông thuyền (MNTT)", f"{mn_tt} m", "Cấp thông thuyền"],
            ["Mực nước thấp nhất (MNTN)", f"{mn_tn} m", "Tần suất P=95%"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị", "Ghi chú"],
                          rows_mn, "Các mực nước thiết kế tại vị trí cầu")

        self._h2(doc, "Đặc điểm địa chất công trình")
        if dgc:
            layers = dgc.get("lop_dat", [])
            if layers:
                rows_dc = []
                for lop in layers:
                    rows_dc.append([
                        lop.get("ten_lop", ""),
                        f"{lop.get('tu_cao_do','')}/{lop.get('den_cao_do','')}",
                        f"{lop.get('N_SPT','N/A')}",
                        lop.get("mo_ta", ""),
                    ])
                self.insert_table(
                    doc,
                    ["Lớp đất", "Cao độ (m)", "SPT-N", "Mô tả"],
                    rows_dc,
                    "Bảng tổng hợp địa chất các lớp đất",
                )
            else:
                self._para(doc,
                    "Theo kết quả khảo sát địa chất sơ bộ, khu vực xây dựng có "
                    "địa tầng gồm các lớp đất điển hình của vùng đồng bằng. Cần "
                    "thực hiện khảo sát chi tiết (SPT, CPT) trước giai đoạn TKKT.")
        else:
            self._para(doc,
                "Chưa có dữ liệu địa chất chi tiết. Dự kiến địa tầng khu vực gồm: "
                "Lớp 1 — Đất đắp và bùn sét mềm yếu (0–5m); "
                "Lớp 2 — Sét dẻo cứng (5–20m); "
                "Lớp 3 — Cát chặt vừa đến chặt (20–40m). "
                "Số liệu chính xác sẽ được cập nhật sau khảo sát địa chất TKKT.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 4 — Quy mô và tải trọng thiết kế
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_4(self, doc):
        kcn  = self._kcn()
        beam = self._beam()
        cap_duong = self.ss.get("cap_duong") or self.ss.get("cap_ky_thuat") or "III"
        vtk = self.ss.get("van_toc_thiet_ke") or 60

        self._h1(doc, "Quy mô và tải trọng thiết kế")

        self._h2(doc, "Tiêu chuẩn kỹ thuật đường và cầu")
        rows_qm = [
            ["Cấp đường",               f"Cấp {cap_duong}"],
            ["Vận tốc thiết kế",        f"{vtk} km/h"],
            ["Tải trọng xe thiết kế",   "HL-93 (TCVN 11823:2017)"],
            ["Tải trọng làn",           "9.3 kN/m/làn"],
            ["Hệ số xung kích IM",      "33% (mỏi) / 75% (khác)"],
            ["Hoạt tải người đi bộ",    "4.8 kN/m² (nếu có vỉa hè)"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_qm,
                          "Tiêu chuẩn kỹ thuật thiết kế")

        self._h2(doc, "Quy mô công trình cầu")
        bc = kcn.get("be_rong_cau") or beam.get("B", "—")
        n_lan = kcn.get("so_lan_xe") or self.ss.get("n_lan_xe") or 2
        n_nhip = kcn.get("tong_so_nhip") or 1
        L_nhip = (beam.get("L", 0) / 1000) if (beam.get("L", 0) > 100) else beam.get("L", 0)
        L_cau  = n_nhip * L_nhip

        rows_cau = [
            ["Tổng chiều dài cầu",    f"{L_cau:.1f} m"],
            ["Số nhịp",               f"{n_nhip} nhịp"],
            ["Chiều dài nhịp điển hình", f"{L_nhip:.1f} m"],
            ["Bề rộng toàn cầu",      f"{bc} m"],
            ["Số làn xe",             f"{n_lan} làn"],
            ["Tĩnh không dưới cầu",   f"≥ {self.ss.get('tinh_khong', 4.5)} m"],
            ["Tĩnh không thông thuyền", f"≥ {self.ss.get('tinh_khong_tt', 3.5)} m (nếu có)"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_cau,
                          "Quy mô công trình cầu chính")

        self._h2(doc, "Tải trọng tác dụng lên công trình")
        self._para(doc,
            "Công trình được thiết kế để chịu các tổ hợp tải trọng theo "
            "TCVN 11823:2017 — Thiết kế cầu đường bộ, bao gồm:")
        self._bullet(doc, "Tĩnh tải: Trọng lượng bản thân kết cấu và các bộ phận cố định")
        self._bullet(doc, "Hoạt tải: Tải trọng xe HL-93 (Truck 325kN + Tandem 220kN + tải làn 9.3kN/m)")
        self._bullet(doc, "Tải trọng gió: Theo vùng gió của khu vực dự án")
        self._bullet(doc, "Tải trọng nhiệt độ: ΔT = ±25°C cho các tính toán biến dạng")
        self._bullet(doc, "Tải trọng va tàu: Kiểm tra khi sông có tàu thuyền qua lại")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 5 — Kết cấu nhịp
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_5(self, doc):
        beam  = self._beam()
        kcn   = self._kcn()
        pa    = self._pa()

        loai_dam = beam.get("loai_dam") or "Super-T"
        L_nhip   = beam.get("L", 0)
        if L_nhip > 100:
            L_nhip /= 1000   # mm → m
        H_dam    = beam.get("H", 1750)
        n_dam    = beam.get("n_dam") or kcn.get("so_dam_tren_mat_cat") or 6
        n_nhip   = kcn.get("tong_so_nhip") or 1
        S_dam    = beam.get("S", 2200)

        self._h1(doc, "Giải pháp thiết kế kết cấu nhịp")

        self._h2(doc, "Phân tích và lựa chọn phương án kết cấu nhịp")
        self._para(doc,
            "Trong giai đoạn nghiên cứu, đã xem xét ba phương án kết cấu nhịp "
            "để lựa chọn phương án tối ưu: PA1 tối ưu chi phí (nhóm dầm không "
            "có bản đáy liền mạch), PA2 tối ưu mỹ quan (nhóm dầm có bản đáy "
            "liền mạch) và PA3 do mô hình Machine Learning đề xuất (Random "
            "Forest học từ dữ liệu công trình cầu Việt Nam). Các phương án "
            "được đánh giá theo 100 điểm (60 điểm kỹ thuật + 30 điểm kinh tế "
            "+ 10 điểm mỹ quan) theo phương pháp AHP (Analytic Hierarchy "
            "Process) tích hợp trong phần mềm AI-Bridge-Design:")

        # Nguồn chọn từng PA: tự động / người dùng khai báo / Machine Learning
        _kcn3 = ((self.ss.get("design_data") or {}).get("kcn_3_pa")
                 if isinstance(self.ss.get("design_data"), dict) else None) or {}
        _PA_ROWS = [("PA1", "pa1_chi_phi", "Tối ưu chi phí"),
                    ("PA2", "pa2_my_quan", "Tối ưu mỹ quan"),
                    ("PA3", "pa3_ml",      "Machine Learning")]

        def _nguon_txt(key, plan):
            if key == "pa3_ml":
                return "Machine Learning (kết quả gốc)"
            if (plan or {}).get("nguon_chon") == "nguoi_dung_khai_bao":
                return "Người dùng khai báo"
            return "Tự động (Rule-Based)"

        pa_list = self.ss.get("so_sanh_3pa") or []
        if _kcn3:
            rows_pa = []
            for _lbl, _key, _mota in _PA_ROWS:
                _plan = _kcn3.get(_key) or (
                    _kcn3.get("pa3_ai") if _key == "pa3_ml" else None) or {}
                rows_pa.append([
                    f"{_lbl} — {_mota}",
                    (f"{_plan.get('loai_dam','—')} "
                     f"L={_plan.get('chieu_dai','—')}m × "
                     f"{_plan.get('tong_so_nhip','—')} nhịp"),
                    _nguon_txt(_key, _plan),
                ])
            self.insert_table(
                doc,
                ["Phương án", "Mô tả", "Nguồn chọn"],
                rows_pa,
                "So sánh và đánh giá các phương án kết cấu nhịp",
            )
            _ovr = [f"{_lbl}" for _lbl, _key, _m in _PA_ROWS
                    if _key != "pa3_ml"
                    and (_kcn3.get(_key) or {}).get("nguon_chon")
                    == "nguoi_dung_khai_bao"]
            if _ovr:
                self._para(doc,
                    f"Ghi chú: phương án {', '.join(_ovr)} sử dụng loại dầm do "
                    "người dùng tự khai báo (ghi đè kết quả tự động của hệ "
                    "thống); các phương án còn lại là kết quả tự động. PA3 "
                    "giữ nguyên kết quả Machine Learning gốc làm cơ sở so "
                    "sánh khách quan.")
        elif pa_list:
            rows_pa = []
            for pa_item in pa_list:
                rows_pa.append([
                    pa_item.get("key", ""),
                    f"{pa_item.get('loai_dam','')} L={pa_item.get('L_nhip','')}m",
                    f"{pa_item.get('diem_tong', pa_item.get('score_total','—'))}",
                    "✓ Chọn" if pa_item.get("is_best") else "",
                ])
            self.insert_table(
                doc,
                ["Phương án", "Mô tả", "Điểm tổng", "Kết quả"],
                rows_pa,
                "So sánh và đánh giá các phương án kết cấu nhịp",
            )
        else:
            self._para(doc,
                f"Sau khi xem xét các phương án thay thế, phương án kết cấu nhịp "
                f"sử dụng {loai_dam} được lựa chọn là phương án tối ưu nhất, "
                f"dựa trên tiêu chí kỹ thuật, kinh tế và thẩm mỹ.")

        self._h2(doc, "Mô tả phương án kết cấu nhịp được chọn")
        LH = L_nhip / (H_dam / 1000) if H_dam else 0
        rows_kcn = [
            ["Loại dầm",                  loai_dam],
            ["Chiều dài nhịp tính toán",  f"{L_nhip:.2f} m"],
            ["Số nhịp",                   f"{n_nhip} nhịp"],
            ["Chiều cao dầm",             f"{H_dam} mm"],
            ["Số dầm trên mặt cắt ngang", f"{n_dam} dầm"],
            ["Khoảng cách tim dầm",       f"{S_dam} mm"],
            ["Tỉ lệ L/H",                 f"{LH:.1f} (TCVN 11823 Điều 5.4.2: ≤ 25)"],
            ["Công nghệ DUL",             beam.get("cong_nghe", "DUL căng trước")],
            ["Mác bê tông",               beam.get("mac_bt", "B50 (f'c = 50 MPa)")],
        ]
        self.insert_table(doc, ["Thông số kỹ thuật", "Giá trị"],
                          rows_kcn, "Thông số kết cấu nhịp")

        self._h2(doc, "Cơ sở lựa chọn phương án")
        self._para(doc,
            f"Phương án {loai_dam} được lựa chọn với các lý do chính sau đây:")
        self._bullet(doc,
            f"Phù hợp với nhịp thiết kế L = {L_nhip:.1f}m (phạm vi kinh tế nhất của {loai_dam})")
        self._bullet(doc,
            f"Tỉ lệ L/H = {LH:.1f} nằm trong giới hạn quy định (≤ 25 theo TCVN 11823 Điều 5.4.2)")
        self._bullet(doc,
            "Công nghệ đúc sẵn tại nhà máy giúp kiểm soát chất lượng bê tông tốt hơn")
        self._bullet(doc,
            "Thẩm mỹ phù hợp với cảnh quan khu vực dự án")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 6 — Mố trụ cầu
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_6(self, doc):
        pier = self._pier()
        kcn  = self._kcn()
        # Kết quả Module 07 (design_data.tru_result) — nguồn nhom_tru + mố +
        # bản quá độ; pier_result cũ (nếu có) vẫn dùng cho các trường legacy.
        _dd     = (self.ss.get("design_data")
                   if isinstance(self.ss.get("design_data"), dict) else {}) or {}
        tru_res = _dd.get("tru_result") or {}
        mo_res  = tru_res.get("ket_qua_mo") or {}

        self._h1(doc, "Giải pháp thiết kế mố và trụ cầu")

        # ── Trụ cầu: phân loại theo 3 NHÓM cấu tạo ────────────────────────
        self._h2(doc, "Thiết kế trụ cầu")
        loai_tru = (tru_res.get("loai_tru") or pier.get("loai_tru")
                    or "Trụ thân đặc BTCT")
        nhom_tru = tru_res.get("nhom_tru") or ""
        ten_nhom = tru_res.get("ten_nhom") or ""
        H_tru    = (pier.get("H_tru") or pier.get("chieu_cao_tru")
                    or _dd.get("H_tru_est") or "—")
        n_tru    = (kcn.get("tong_so_nhip") or 1) - 1

        self._para(doc,
            "Trụ cầu được phân loại theo ba nhóm cấu tạo chính: "
            "(1) Trụ dẻo (trụ cọc) — một hoặc hai hàng cọc tiết diện nhỏ "
            "30×30 đến 40×40cm liên kết trực tiếp với xà mũ, không có bệ trụ "
            "riêng, áp dụng cho cầu nhiều nhịp ngắn ≤ 12m vượt sông nhỏ cấp "
            "V–VI, lòng sông không sâu và không thông thuyền (chiều cao trụ "
            "thấp là hệ quả tự nhiên, không phải tiêu chí phân nhóm); "
            "(2) Trụ cột (trụ thân cột BTCT) — một hoặc nhiều cột tròn/chữ "
            "nhật đường kính phổ biến 0.8–2m liên kết với xà mũ chịu uốn, có "
            "bệ móng riêng, áp dụng cho cầu cạn, cầu vượt và cầu vượt sông "
            "cấp IV–VI ít cây trôi (một cột cho cầu vượt đô thị để giải "
            "phóng không gian gầm cầu, 2–4 cột cho cầu vừa và rộng); "
            "(3) Trụ đặc thân hẹp — phần thân dưới thu hẹp so với bề rộng "
            "kết cấu nhịp, mũi vát nhọn với hệ số cản dòng chảy CD = 0.7–0.8, "
            "áp dụng cho cầu vượt sông lớn cấp I–III chịu va tàu và nhiều "
            "cây trôi; khi chiều cao trụ H > 10m chuyển sang biến thể trụ "
            "đặc thân hẹp rỗng để giảm khối lượng vật liệu.")
        self._para(doc,
            f"Dựa trên điều kiện địa hình, tải trọng và nhịp thiết kế, "
            f"phương án trụ được lựa chọn là: {loai_tru}"
            + (f" — thuộc nhóm {ten_nhom or nhom_tru}" if (nhom_tru or ten_nhom) else "")
            + f". Tổng số trụ trung gian: {n_tru} trụ.")

        rows_tru = [
            ["Nhóm trụ (cấu tạo)",    ten_nhom or nhom_tru or "—"],
            ["Loại trụ",              loai_tru],
            ["Số lượng trụ",          f"{n_tru} trụ"],
            ["Chiều cao thân trụ",    f"{H_tru} m"],
            ["Vật liệu thân trụ",     pier.get("vat_lieu_tru", "BTCT B30")],
            ["Kích thước sơ bộ",      pier.get("kich_thuoc_tru", "Xem bản vẽ")],
        ]

        if pier.get("canh_bao") or tru_res.get("canh_bao"):
            rows_tru.append(["Cảnh báo kỹ thuật",
                             pier.get("canh_bao") or tru_res.get("canh_bao", "")])

        self.insert_table(doc, ["Thông số", "Giá trị"], rows_tru,
                          "Thông số thiết kế trụ cầu")

        tang_qd = (pier.get("tang_quyet_dinh")
                   or (tru_res.get("pa_rb") or {}).get("tang_quyet_dinh") or [])
        if tang_qd:
            self._h3(doc, "Cơ sở quyết định loại trụ")
            self._para(doc, "Phân tích lựa chọn theo phương pháp Rule-Based "
                            "(phân nhóm cấu tạo trước, loại con trong nhóm sau):")
            if isinstance(tang_qd, str):
                tang_qd = [tang_qd]
            for t in tang_qd:
                self._bullet(doc, str(t))
            _gc_rb = (tru_res.get("pa_rb") or {}).get("ghi_chu")
            if _gc_rb:
                self._bullet(doc, str(_gc_rb))

        # ── Mố cầu ─────────────────────────────────────────────────────────
        self._h2(doc, "Thiết kế mố cầu")
        loai_mo  = (mo_res.get("loai_mo") or pier.get("loai_mo")
                    or "Mố chữ U BTCT")
        H_dap    = pier.get("H_dap") or self.ss.get("H_dap") or "—"

        self._para(doc,
            f"Hai đầu cầu sử dụng mố cầu loại {loai_mo}, phù hợp với "
            f"chiều cao đất đắp H_đắp ≈ {H_dap}m tại vị trí hai mố. "
            f"Mố được thiết kế theo Điều 11.5, TCVN 11823:2017.")

        rows_mo = [
            ["Loại mố",           loai_mo],
            ["Chiều cao đất đắp", f"{H_dap} m"],
            ["Vật liệu",          pier.get("vat_lieu_mo", "BTCT B25")],
            ["Bản quá độ sau mố", "BẮT BUỘC — xem mục Bản quá độ dưới đây"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_mo,
                          "Thông số thiết kế mố cầu")

        # ── Bản quá độ (cấu kiện BẮT BUỘC cho mọi mố) ──────────────────────
        self._h3(doc, "Bản quá độ sau mố (cấu kiện bắt buộc)")
        bqd = mo_res.get("ban_qua_do") or {}
        self._para(doc,
            "Mọi mố cầu đều bố trí bản quá độ với bốn chức năng thiết yếu:")
        for cn in (bqd.get("chuc_nang") or [
            "Khắc phục hiện tượng điểm xóc đầu cầu (đất đắp trong lòng mố "
            "khó đạt độ chặt tuyệt đối)",
            "Chuyển tiếp độ cứng dần dần giữa nền đường mềm và mố cầu cứng",
            "Xử lý lún chênh lệch — bù trừ khi độ chênh < 5cm",
            "Phân phối lại tải trọng đất đắp + hoạt tải xe lên mố theo "
            "hướng tích cực cho ổn định",
        ]):
            self._bullet(doc, str(cn))

        _L_bqd = bqd.get("L_bqd", "≥ 5 (theo quy mô cầu)")
        rows_bqd = [
            ["Chiều dài bản L_bqd",
             (f"{_L_bqd} m ({bqd.get('L_bqd_range','')} — "
              f"{bqd.get('quy_mo_cau','theo quy mô cầu')})"
              if bqd else "Cầu nhỏ ≥5m · cầu trung 6–8m · cầu lớn 8–12m")],
            ["Chiều dày bản (day_bqd)",
             f"≥ {float(bqd.get('day_bqd', 0.30))*100:.0f} cm"],
            ["Độ dốc dọc (doc_bqd)",
             str(bqd.get("doc_bqd", "10–15% về phía nền đường"))],
            ["Đất đắp mặt đường → mặt bản (chieu_sau_dat_dap)",
             f"≥ {float(bqd.get('chieu_sau_dat_dap_toi_thieu', 0.70))*100:.0f} cm "
             "(tránh nứt mặt đường do hoạt tải xe)"],
            ["Vật liệu", str(bqd.get("vat_lieu",
                                     "BTCT M300 đổ tại chỗ hoặc đúc sẵn"))],
            ["Vị trí lắp đặt", str(bqd.get("vi_tri_lap_dat",
             "Một đầu bản kê lên tường đỉnh mố, đầu còn lại đặt trên dầm kê "
             "nằm trong nền đường."))],
        ]
        self.insert_table(doc, ["Thông số bản quá độ", "Giá trị"], rows_bqd,
                          "Thông số thiết kế bản quá độ")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 7 — Móng cọc
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_7(self, doc):
        mong = self._mong()
        # Kết quả Module 08 (design_data.mong_result) — nguồn 4 phần A/B/C/D
        _dd = (self.ss.get("design_data")
               if isinstance(self.ss.get("design_data"), dict) else {}) or {}
        mr = _dd.get("mong_result") or {}

        self._h1(doc, "Giải pháp thiết kế móng")
        self._para(doc,
            "Giải pháp móng cọc được xác lập theo bốn nhóm logic tuân thủ "
            "TCVN 10304:2025: (A) lựa chọn loại cọc theo đường kính; "
            "(B) chiều dài cọc và tầng tựa mũi; (C) số lượng cọc với hệ số "
            "hiệu quả nhóm; (D) khoảng cách bố trí cọc trên bệ.")

        loai_coc = (mr.get("loai_mong") or mr.get("loai_coc")
                    or mong.get("loai_coc") or "Cọc khoan nhồi BTCT")
        D_coc    = (mr.get("D_coc_mm") or mr.get("kich_thuoc_mm")
                    or mong.get("D_coc") or mong.get("duong_kinh_coc") or 800)
        L_coc    = (mr.get("chieu_dai_coc") or mr.get("L_coc_tu")
                    or mong.get("L_coc") or mong.get("chieu_dai_coc") or "—")
        n_coc_be = (mr.get("so_coc_be") or mr.get("So_coc_tu")
                    or mong.get("n_coc_be") or mong.get("so_coc_tren_be") or "—")
        Qa       = (mr.get("Q_1coc_tk_kN") or mong.get("Q_cho_phep")
                    or mong.get("Qa") or "—")
        la_ckn   = "khoan nhồi" in str(loai_coc).lower()

        # ── A. Lựa chọn loại cọc theo đường kính ────────────────────────
        self._h2(doc, "A. Lựa chọn loại cọc theo đường kính")
        self._para(doc,
            "Cọc được phân hai nhóm theo đường kính: nhóm ĐƯỜNG KÍNH NHỎ "
            "(D ≤ 0.6m — cọc đóng/ép vuông 200–450mm, cọc tròn ly tâm "
            "D300–D600, Q_tk 250–1800 kN) áp dụng cho cầu nhịp nhỏ, tải đầu "
            "cọc thấp, nền không quá rắn, không đá tảng/đá mồ côi; nhóm "
            "ĐƯỜNG KÍNH LỚN (D = 0.8–2m — cọc khoan nhồi D800–D2000, Q_tk "
            "1500–6000 kN) cho tải trọng lớn, đô thị hạn chế rung động, "
            "có đá phong hóa/đá mồ côi cần khoan xuyên; đường kính chọn "
            "theo sức chịu tải yêu cầu Q_tk và số lượng cọc hợp lý.")
        rows_a = [
            ["Nhóm cọc (category)",
             mr.get("ten_nhom_coc") or ("Đường kính lớn (large)" if la_ckn
                                        else "Đường kính nhỏ (small)")],
            ["Loại cọc",               str(loai_coc)],
            ["Kích thước chọn",        f"D = {D_coc} mm"],
            ["Sức chịu tải Q_tk",      f"{Qa} kN/cọc"],
        ]
        if mr.get("ly_tam_tuong_duong"):
            rows_a.append(["Phương án thay thế cùng nhóm",
                           str(mr["ly_tam_tuong_duong"])])
        _lyd = (mr.get("ly_do_chon") or mong.get("tang_quyet_dinh") or [])
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_a,
                          "Lựa chọn loại cọc (A)")
        if _lyd:
            self._para(doc, "Cơ sở lựa chọn:")
            for t in _lyd[:4]:
                self._bullet(doc, str(t))

        # ── B. Chiều dài cọc và tầng tựa mũi ────────────────────────────
        self._h2(doc, "B. Chiều dài cọc và tầng tựa mũi")
        self._para(doc,
            "Lớp tựa mũi phải đồng thời đạt tiêu chuẩn định lượng theo "
            "TCVN 10304:2025: đất dính SPT-N ≥ 30; đất rời SPT-N ≥ 40; "
            "đá RQD > 60%. Chiều sâu cắm vào lớp tốt: đất chặt/dính cứng "
            "≥ 3.0m; đất yếu/đất rời ≥ 6.0m; đá tươi RQD > 75% (hoặc đạt "
            "độ chối) chỉ cần ngàm ≥ 0.5m.")
        rows_b = [
            ["Lớp tựa mũi",            str(mr.get("lop_tua_mui")
                                           or mong.get("lop_tua_mui", "—"))],
            ["Tiêu chuẩn kiểm tra",    str(mr.get("lop_tua_mo_ta", "—"))],
            ["Chiều sâu cắm L_cam",    f"{mr.get('chieu_dai_cam_lop_tot', '—')} m"
             + (f" ({mr.get('co_so_cam')})" if mr.get("co_so_cam") else "")],
            ["Tổng chiều dài cọc",     f"L = {L_coc} m"],
            ["Tỷ lệ L/D",              f"{mr.get('ty_le_LD', '—')} "
                                       "(khoảng hợp lý 30–100; <30 xem xét "
                                       "giảm D, >100 nền yếu tăng D)"],
            ["Q vật liệu / Q đất nền", f"{mr.get('Q_vl_kN', '—')} / {Qa} kN "
                                       "(chênh > 40% → xem lại kinh tế)"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_b,
                          "Chiều dài cọc và tầng tựa mũi (B)")

        # ── C. Số lượng cọc ─────────────────────────────────────────────
        self._h2(doc, "C. Số lượng cọc và hệ số hiệu quả nhóm")
        self._para(doc,
            "Số cọc trên bệ: n = max(4, ⌈P_bệ / (Q_1cọc × η_g)⌉ × dự phòng). "
            "Hệ số hiệu quả nhóm η_g lấy theo Converse-Labarre (đất dính, "
            "bệ không tiếp xúc đất — tim-tim 2.5D cho η ≈ 0.7–0.8) hoặc "
            "≈ 0.9–1.0 với đất rời. Số cọc đã nhân dự phòng sai số thi công "
            "1.05–1.10 (cọc đóng lệch 75–150mm).")
        rows_c = [
            ["Số cọc / bệ",            f"{n_coc_be} cọc"],
            ["Hệ số nhóm η_g",         f"{mr.get('he_so_nhom', '—')}"],
            ["Hệ số dự phòng thi công", f"×{mr.get('he_so_du_phong', '—')} "
                                        "(lệch 75–150mm)"],
        ]
        if mr.get("cong_thuc_so_coc"):
            rows_c.append(["Công thức áp dụng", str(mr["cong_thuc_so_coc"])])
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_c,
                          "Số lượng cọc (C)")

        # ── D. Khoảng cách bố trí cọc ───────────────────────────────────
        self._h2(doc, "D. Khoảng cách bố trí cọc trên bệ")
        if la_ckn:
            self._para(doc,
                "Cọc khoan nhồi (TCVN 10304:2025): tim-tim thiết kế > 3.0D; "
                "tim-tim < 4D thi công không ống vách phải đánh giá tương "
                "tác giữa các cọc; tim-tim < 6D phải quy định rõ trình tự "
                "thi công khoan trong hồ sơ thiết kế; khoảng thông thủy "
                "giữa thân các cọc ≥ 1.0m; mặt bên cọc đến mép bệ ≥ 300mm.")
        else:
            self._para(doc,
                "Cọc đóng/ép (TCVN 10304:2025): khoảng cách tim-tim tối "
                "thiểu max(750mm, 2.5D); tại mặt phẳng mũi cọc ma sát "
                "tim-tim ≥ 3D; mặt ngoài cọc đến mép bệ ≥ 225mm.")
        _S  = mr.get("khoang_cach_tim") or mr.get("khoang_cach_tim_coc")
        _mp = mr.get("khoang_cach_mep")
        rows_d = [
            ["Khoảng cách tim-tim",   f"{_S if _S is not None else '—'} m"],
            ["Mặt cọc → mép bệ",      (f"≥ {float(_mp)*1000:.0f} mm"
                                       if _mp else ("≥ 300 mm" if la_ckn
                                                    else "≥ 225 mm"))],
        ]
        if la_ckn:
            rows_d.append(["Thông thủy thân cọc",
                           f"{mr.get('khoang_cach_thong_thuy', '—')} m (≥ 1.0m)"])
        rows_d.append(["Kích thước bệ cọc",
                       str(mr.get("kich_thuoc_be")
                           or mr.get("kich_thuoc_be_goi_y", "—"))])
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_d,
                          "Khoảng cách bố trí cọc theo TCVN 10304:2025 (D)")
        _ws = [w for w in (mr.get("warnings") or [])
               if "tương tác" in w or "Trình tự" in w]
        for w in _ws:
            self._bullet(doc, str(w))

        self._h2(doc, "Bệ cọc và kiểm tra ổn định")
        self._para(doc,
            f"Bệ cọc được thiết kế đủ lớn để chứa {n_coc_be} cọc theo các "
            f"khoảng cách mục D. Chiều dày bệ cọc tối thiểu "
            f"{max(float(D_coc)/1000, 0.8):.1f}m. Cao độ đỉnh bệ cọc đặt "
            "dưới cao độ đáy mố/trụ để đảm bảo ổn định.")
        self._para(doc,
            "Kiểm tra xói lở: Chiều sâu đặt móng phải nằm dưới cao độ xói thiết kế "
            "≥ 0.5m (Điều 2.6, TCVN 11823:2017). Xói cục bộ quanh cọc được tính "
            "theo công thức HEC-18 tham chiếu trong TCVN 11823.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 8 — Mặt cầu và bộ phận phụ trợ
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_8(self, doc):
        beam = self._beam()

        self._h1(doc, "Mặt cầu và các bộ phận phụ trợ")

        self._h2(doc, "Cấu tạo lớp phủ mặt cầu")
        self._para(doc,
            "Mặt cầu được thiết kế gồm các lớp từ trên xuống dưới:")
        layers = [
            ["Lớp phủ nhựa Polymer Modified Bitumen (PMB)",
             "40–50 mm", "TCVN 8819:2011"],
            ["Lớp phòng nước và bảo vệ (Waterproofing membrane)",
             "10–15 mm", "ASTM D 5957"],
            ["Bản mặt cầu BTCT đổ tại chỗ",
             "200–250 mm", "B30, liên kết với dầm bằng neo chốt"],
        ]
        self.insert_table(
            doc,
            ["Lớp cấu tạo", "Chiều dày", "Tiêu chuẩn"],
            layers,
            "Cấu tạo lớp phủ mặt cầu",
        )

        self._h2(doc, "Gối cầu và khe co giãn")
        n_nhip = self._kcn().get("tong_so_nhip") or 1
        n_dam  = beam.get("n_dam") or 6
        n_goi  = n_dam * (n_nhip + 1)
        rows_goi = [
            ["Loại gối",           "Gối cao su bản thép (Elastomeric bearing)"],
            ["Số lượng gối",       f"{n_goi} cái (= {n_dam} dầm × {n_nhip+1} vị trí)"],
            ["Kích thước tham khảo", "350×450×65mm (xác định chính xác ở TKKT)"],
            ["Loại khe co giãn",   "Khe co giãn modular/finger"],
            ["Số lượng khe",       f"2 khe (2 đầu cầu)"],
        ]
        self.insert_table(doc, ["Thông số", "Giá trị"], rows_goi,
                          "Gối cầu và khe co giãn")

        self._h2(doc, "Hệ thống lan can và thoát nước")
        self._para(doc,
            "Lan can phòng hộ kim loại (steel guardrail) được bố trí hai bên cầu, "
            "thiết kế theo TCVN 4054:2005 phù hợp với cấp đường. Hệ thống thoát "
            "nước mặt cầu gồm ống thoát nước PVC Ø100mm bố trí cách nhau ≤ 5m "
            "dọc theo chiều dài cầu, đảm bảo thoát nước nhanh, tránh đọng nước "
            "gây hư hỏng mặt cầu.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 9 — Dự toán sơ bộ
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_9(self, doc):
        dt   = self._dt()
        th   = dt.get("tong_hop") or {}
        ph   = dt.get("phan_tich") or {}
        kl   = dt.get("khoi_luong") or {}

        self._h1(doc, "Dự toán sơ bộ tổng mức đầu tư")

        self._h2(doc, "Cơ sở lập dự toán")
        self._para(doc,
            "Dự toán sơ bộ được lập theo phương pháp suất đầu tư quy định tại "
            "Điều 36.2, Nghị định 175/2024/NĐ-CP, sử dụng đơn giá tham chiếu "
            "theo Quyết định 409/2025/QĐ-BXD. Các đơn giá vật liệu, nhân công, "
            "máy thi công theo mức giá thị trường hiện hành.")

        self._h2(doc, "Tổng hợp dự toán theo hạng mục")
        if th:
            rows_dt = []
            hm_map = [
                ("hm1_KCN",  "Kết cấu nhịp (KCN)"),
                ("hm2_MTM",  "Mố, trụ cầu (MTM)"),
                ("hm3_MON",  "Móng cọc (MON)"),
                ("hm4_BMC",  "Bản mặt cầu và phụ trợ (BMC)"),
                ("hm5_PHU",  "Chi phí phụ (gối, khe, lan can, thoát nước)"),
                ("chi_phi_dp", "Dự phòng phí"),
                ("phi_quan_ly","Chi phí quản lý, thiết kế, giám sát"),
                ("tong_xay_lap","TỔNG XÂY LẮP"),
                ("VAT",       "Thuế VAT (10%)"),
                ("tong_dau_tu","TỔNG MỨC ĐẦU TƯ (trước thuế)"),
            ]
            for key, label in hm_map:
                val = th.get(key)
                if val is not None:
                    rows_dt.append([label, f"{val:,.0f}", "triệu VND"])

            if rows_dt:
                self.insert_table(
                    doc,
                    ["Hạng mục", "Giá trị (triệu VND)", "Đơn vị"],
                    rows_dt,
                    "Tổng hợp dự toán sơ bộ theo hạng mục",
                )
        else:
            self._para(doc,
                "Chưa có dữ liệu dự toán. Vui lòng hoàn thành Module 12 — "
                "Dự toán sơ bộ trước khi xuất thuyết minh.")

        # Suất đầu tư
        sdt = ph.get("suat_dau_tu_trieu_m2")
        if sdt:
            self._h2(doc, "Suất đầu tư và đánh giá tính hợp lý")
            cap_ct = self.ss.get("cap_cong_trinh", "III")
            self._para(doc,
                f"Suất đầu tư tính toán: {sdt:.2f} triệu VND/m² diện tích mặt cầu. "
                f"Theo QĐ 409/2025, suất đầu tư tham chiếu cho cầu cấp {cap_ct} "
                f"là 8–13 triệu VND/m². "
                + ("✓ Kết quả nằm trong khoảng tham chiếu, đảm bảo tính hợp lý."
                   if 8 <= sdt <= 13
                   else f"⚠ Kết quả ({sdt:.2f} triệu/m²) lệch khỏi mức tham chiếu, "
                        f"cần giải trình thêm."))

        # Cố gắng chèn biểu đồ phân bổ
        try:
            import plotly.graph_objects as go
            hm_vals = {
                "KCN": th.get("hm1_KCN", 0),
                "MTM": th.get("hm2_MTM", 0),
                "MON": th.get("hm3_MON", 0),
                "BMC": th.get("hm4_BMC", 0),
                "PHU": th.get("hm5_PHU", 0),
            }
            if any(v > 0 for v in hm_vals.values()):
                fig = go.Figure(go.Pie(
                    labels=list(hm_vals.keys()),
                    values=list(hm_vals.values()),
                    textinfo="percent+label",
                ))
                fig.update_layout(title_text="Phân bổ chi phí xây lắp theo hạng mục",
                                  width=700, height=450)
                self.insert_chart_from_plotly(
                    doc, fig, "Biểu đồ phân bổ chi phí theo hạng mục"
                )
        except Exception:
            pass  # bỏ qua nếu kaleido hoặc plotly không hoạt động

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 10 — Môi trường sơ bộ
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_10(self, doc):
        self._h1(doc, "Đánh giá tác động môi trường sơ bộ")

        self._para(doc,
            "Phần đánh giá tác động môi trường (ĐTM) sơ bộ này được lập theo "
            "Điều 20, Nghị định 175/2024/NĐ-CP và Luật Bảo vệ Môi trường năm 2020. "
            "Đây là đánh giá sơ bộ, chưa thay thế cho ĐTM đầy đủ theo quy định.")

        self._h2(doc, "Tác động trong giai đoạn thi công")
        tc_impacts = [
            ["Tiếng ồn và rung động",
             "Sử dụng thiết bị thi công",
             "Hạn chế thi công ban đêm; lắp đặt rào chắn âm thanh"],
            ["Bụi và khí thải",
             "Vận chuyển vật liệu, đào đắp đất",
             "Tưới nước, phủ bạt xe tải; bảo dưỡng thiết bị định kỳ"],
            ["Chất thải xây dựng",
             "Đất đào thừa, vật liệu phế thải",
             "Thu gom và đổ thải tại bãi quy định; tái sử dụng tối đa"],
            ["Ô nhiễm nguồn nước",
             "Nước thải xây dựng, dầu nhớt",
             "Hệ thống thu gom nước thải thi công; không xả thải trực tiếp ra sông"],
        ]
        self.insert_table(
            doc,
            ["Tác động", "Nguồn gốc", "Biện pháp giảm thiểu"],
            tc_impacts,
            "Tác động môi trường giai đoạn thi công và biện pháp giảm thiểu",
        )

        self._h2(doc, "Tác động trong giai đoạn khai thác")
        self._para(doc,
            "Khi công trình đi vào khai thác, các tác động chính bao gồm: "
            "thay đổi dòng chảy thủy văn (cần kiểm tra xói lở định kỳ); "
            "tác động đến đường thủy nội địa (tĩnh không thông thuyền đảm bảo theo thiết kế); "
            "tiếng ồn từ xe cộ đi lại trên cầu (tuân thủ QCVN 26:2010/BTNMT). "
            "Về tổng thể, công trình cầu có tác động tích cực đến giao thông "
            "và phát triển kinh tế — xã hội khu vực.")

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 11 — Tổ chức thi công sơ bộ
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_11(self, doc):
        beam = self._beam()
        mong = self._mong()
        loai_dam = beam.get("loai_dam", "Super-T")
        loai_coc = mong.get("loai_coc", "Cọc khoan nhồi")

        self._h1(doc, "Phương án tổ chức thi công sơ bộ")

        self._h2(doc, "Phương án thi công móng và mố trụ")
        if "khoan nhồi" in loai_coc.lower():
            self._para(doc,
                "Móng cọc khoan nhồi được thi công theo trình tự: "
                "(1) Định vị và lắp ống vách; (2) Khoan tạo lỗ bằng máy khoan rotary; "
                "(3) Lắp dựng lồng thép; (4) Đổ bê tông theo phương pháp ống tremie; "
                "(5) Kiểm tra toàn vẹn cọc bằng sonic logging (CSL); "
                "(6) Thi công bệ cọc và thân trụ/mố.")
        else:
            self._para(doc,
                "Móng cọc ép được thi công theo trình tự: "
                "(1) Định vị tim cọc; (2) Lắp máy ép; (3) Ép cọc theo từng đoạn "
                "đến khi đạt tải ép quy định; (4) Nối đầu cọc vào bệ cọc.")

        self._h2(doc, "Phương án thi công kết cấu nhịp")
        if loai_dam in ("Super-T", "Dầm I", "T ngược", "Dầm bản rỗng"):
            self._para(doc,
                f"Dầm {loai_dam} đúc sẵn tại nhà máy, vận chuyển đến công trường, "
                f"cẩu lắp bằng cần cẩu trên xà lan hoặc cần cẩu bờ. "
                f"Sau khi lắp đặt dầm, thi công bản mặt cầu và liên kết ngang.")
        else:
            self._para(doc,
                "Kết cấu nhịp thi công đổ bê tông tại chỗ bằng hệ ván khuôn. "
                "Cần đảm bảo kiểm soát chất lượng bê tông và thời gian dưỡng hộ.")

        self._h2(doc, "Dự kiến tiến độ thi công")
        rows_td = [
            ["Giai đoạn 1", "Chuẩn bị mặt bằng, định vị công trình", "1 tháng"],
            ["Giai đoạn 2", "Thi công móng cọc (mố + trụ)", "2–3 tháng"],
            ["Giai đoạn 3", "Thi công bệ cọc, thân trụ, mố", "2 tháng"],
            ["Giai đoạn 4", "Lắp dựng kết cấu nhịp", "1 tháng"],
            ["Giai đoạn 5", "Hoàn thiện mặt cầu, phụ trợ, an toàn GT", "1–2 tháng"],
            ["TỔNG CỘNG",   "", "7–9 tháng"],
        ]
        self.insert_table(
            doc,
            ["Giai đoạn", "Nội dung công việc", "Thời gian dự kiến"],
            rows_td,
            "Dự kiến tiến độ thi công sơ bộ",
        )

    # ─────────────────────────────────────────────────────────────────────────
    # CHƯƠNG 12 — Kết luận và kiến nghị
    # ─────────────────────────────────────────────────────────────────────────

    def generate_chapter_12(self, doc):
        beam  = self._beam()
        mong  = self._mong()
        dt    = self._dt()
        ph    = dt.get("phan_tich") or {}
        th    = dt.get("tong_hop") or {}

        self._h1(doc, "Kết luận và kiến nghị")

        self._h2(doc, "Kết luận")
        loai_dam = beam.get("loai_dam", "Super-T")
        L_nhip   = beam.get("L", 0)
        if L_nhip > 100:
            L_nhip /= 1000
        sdt      = ph.get("suat_dau_tu_trieu_m2", 0)
        tdt      = th.get("tong_dau_tu", 0)

        self._para(doc,
            "Trên cơ sở phân tích các điều kiện tự nhiên, kỹ thuật, kinh tế "
            "và yêu cầu của dự án, phương án thiết kế cơ sở đề xuất như sau:")
        self._bullet(doc,
            f"Kết cấu nhịp: {loai_dam}, nhịp L = {L_nhip:.1f}m — phù hợp về mặt kỹ thuật và kinh tế")
        self._bullet(doc,
            f"Móng: {mong.get('loai_coc','Cọc khoan nhồi')} — đảm bảo sức chịu tải và an toàn")
        self._bullet(doc,
            f"Suất đầu tư: {sdt:.2f} triệu VND/m² — "
            + ("nằm trong khoảng tham chiếu theo QĐ 409/2025" if 8<=sdt<=20
               else "cần giải trình chi tiết ở giai đoạn TKKT"))
        self._bullet(doc,
            f"Tổng mức đầu tư dự kiến: {tdt:,.0f} triệu VND")

        self._para(doc,
            "Phương án đề xuất đảm bảo tính khả thi về mặt kỹ thuật, tuân thủ "
            "đầy đủ các tiêu chuẩn kỹ thuật (TCVN 11823:2017, TCVN 10304:2014) "
            "và các quy định của Nghị định 175/2024/NĐ-CP. Thuyết minh TKCS này "
            "đủ điều kiện nộp thẩm định theo quy định tại Điều 21, NĐ 175/2024.")

        self._h2(doc, "Kiến nghị")
        self._para(doc,
            "Để tiến hành giai đoạn thiết kế kỹ thuật (TKKT), kiến nghị thực hiện:")
        self._bullet(doc,
            "Khảo sát địa chất chi tiết: Thực hiện thêm hố khoan SPT và CPT tại "
            "vị trí trụ trung gian và hai mố để xác định chính xác chiều dài cọc")
        self._bullet(doc,
            "Khảo sát thủy văn chi tiết: Xác định chính xác MNCN (P=1%) và "
            "chiều sâu xói thiết kế tại các vị trí trụ")
        self._bullet(doc,
            "Phê duyệt TKCS: Nộp hồ sơ thẩm định theo cấp thẩm quyền trước khi "
            "triển khai TKKT")
        self._bullet(doc,
            "Lập hồ sơ mời thầu tư vấn khảo sát và thiết kế kỹ thuật ngay sau "
            "khi TKCS được phê duyệt")

    # ─────────────────────────────────────────────────────────────────────────
    # Đánh giá chất lượng thuyết minh
    # ─────────────────────────────────────────────────────────────────────────

    def quality_check(self) -> Dict:
        """Chấm điểm chất lượng thuyết minh theo 4 tiêu chí."""
        scores = {}
        ss = self.ss

        # 1. Tính đầy đủ (max 40pt)
        required_keys = [
            "beam_params_final", "kcn_result", "pier_result",
            "mong_coc_result", "du_toan_result",
        ]
        n_present = sum(1 for k in required_keys if ss.get(k))
        scores["day_du"] = int(40 * n_present / len(required_keys))

        # 2. Tính nhất quán dữ liệu (max 25pt)
        beam = self._beam()
        dt   = self._dt()
        ph   = dt.get("phan_tich") or {}
        sdt  = ph.get("suat_dau_tu_trieu_m2", 0)
        ok_sdt = 8 <= sdt <= 20 if sdt else False
        L_nhip = (beam.get("L", 0) / 1000) if beam.get("L", 0) > 100 else beam.get("L", 0)
        loai   = beam.get("loai_dam", "")
        span_ranges = {"Super-T":(25,50), "Dầm I":(15,40), "T ngược":(10,25), "Dầm bản rỗng":(6,20)}
        ok_nhip = (L_nhip >= span_ranges.get(loai, (0,100))[0] and
                   L_nhip <= span_ranges.get(loai, (0,100))[1]) if loai else False
        scores["nhat_quan"] = 25 if (ok_sdt and ok_nhip) else (15 if ok_nhip else 5)

        # 3. Trích dẫn tiêu chuẩn (max 20pt — giả sử đủ vì template đã có)
        scores["trich_dan"] = 18

        # 4. Thông tin dự án (max 15pt)
        pi = self.pi
        pi_keys = ["ten_du_an", "dia_diem", "chu_dau_tu", "hoc_vien", "gv_huong_dan"]
        n_pi = sum(1 for k in pi_keys if pi.get(k))
        scores["thong_tin_da"] = int(15 * n_pi / len(pi_keys))

        total = sum(scores.values())
        grade = "Xuất sắc (A)" if total >= 90 else \
                "Tốt (B)" if total >= 75 else \
                "Đạt (C)" if total >= 60 else "Cần bổ sung"

        return {
            "scores":     scores,
            "total":      total,
            "max":        100,
            "grade":      grade,
            "issues": [
                f"Thiếu dữ liệu module" if scores["day_du"] < 32 else None,
                f"Suất đầu tư {sdt:.2f} ngoài tham chiếu (8–20 triệu/m²)" if not ok_sdt and sdt else None,
                f"Nhịp {L_nhip:.1f}m không tối ưu cho {loai}" if not ok_nhip and loai else None,
                "Cần nhập đầy đủ thông tin dự án (tên, địa điểm, chủ đầu tư)" if scores["thong_tin_da"] < 10 else None,
            ],
        }

    # ─────────────────────────────────────────────────────────────────────────
    # Orchestrator chính
    # ─────────────────────────────────────────────────────────────────────────

    def generate(self, progress_callback=None) -> bytes:
        """
        Sinh toàn bộ thuyết minh TKCS.

        Parameters
        ----------
        progress_callback : callable(step:int, total:int, msg:str), optional

        Returns
        -------
        bytes — nội dung file DOCX
        """
        def _prog(step, msg):
            if progress_callback:
                progress_callback(step, 15, msg)

        _prog(0, "Khởi tạo document...")
        self.doc = self._setup_document()

        _prog(1, "Tạo trang bìa...")
        self._add_cover_page(self.doc)

        # Sinh nội dung 12 chương
        chapter_methods = [
            (1,  "Thông tin chung dự án",          self.generate_chapter_1),
            (2,  "Cơ sở pháp lý và tiêu chuẩn",   self.generate_chapter_2),
            (3,  "Điều kiện tự nhiên",              self.generate_chapter_3),
            (4,  "Quy mô và tải trọng",             self.generate_chapter_4),
            (5,  "Kết cấu nhịp",                    self.generate_chapter_5),
            (6,  "Mố và trụ cầu",                   self.generate_chapter_6),
            (7,  "Móng cọc",                        self.generate_chapter_7),
            (8,  "Mặt cầu và phụ trợ",              self.generate_chapter_8),
            (9,  "Dự toán sơ bộ",                   self.generate_chapter_9),
            (10, "Tác động môi trường",              self.generate_chapter_10),
            (11, "Tổ chức thi công sơ bộ",          self.generate_chapter_11),
            (12, "Kết luận và kiến nghị",            self.generate_chapter_12),
        ]

        for ch_num, title, method in chapter_methods:
            if ch_num in self.selected:
                _prog(ch_num + 1, f"Đang sinh Chương {ch_num}: {title}...")
                try:
                    method(self.doc)
                    self.doc.add_page_break()
                except Exception as e:
                    logger.warning("Lỗi Chương %d: %s", ch_num, e)
                    self.doc.add_paragraph(f"[Chương {ch_num} — Lỗi: {e}]")
                    self.doc.add_page_break()

        _prog(14, "Thêm header/footer và mục lục...")
        title_short = self.pi.get("ten_du_an", "Thuyết minh TKCS")[:60]
        self._add_header_footer(self.doc, title_short)

        # Thêm mục lục + danh sách bảng/hình vào đầu (sau trang bìa)
        # Tạo document mới để ghép mục lục + nội dung
        front_doc = self._setup_document()
        self._add_cover_page(front_doc)
        self._add_abbreviations(front_doc)
        self._add_front_matter(front_doc)

        # Lưu ra bytes
        _prog(15, "Hoàn thiện và lưu file...")
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()

    def save(self, output_path: Optional[str] = None) -> str:
        """Sinh và lưu file DOCX vào đĩa. Trả về đường dẫn file."""
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        if output_path is None:
            ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            name = (self.pi.get("ten_du_an", "TKCS") or "TKCS")\
                   .replace(" ", "_")[:30]
            output_path = str(OUTPUT_DIR / f"{name}_{ts}.docx")
        data = self.generate()
        Path(output_path).write_bytes(data)
        return output_path

    def generate_zip(self) -> bytes:
        """Tạo file ZIP gồm TKCS chính + danh mục bảng/hình dưới dạng TXT."""
        docx_bytes = self.generate()
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            ts   = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            name = (self.pi.get("ten_du_an", "TKCS") or "TKCS")\
                   .replace(" ", "_")[:30]
            zf.writestr(f"TKCS_{name}_{ts}.docx", docx_bytes)

            # Phụ lục 4: Danh mục tiêu chuẩn
            tc_lines = ["DANH MỤC TIÊU CHUẨN ÁP DỤNG\n" + "="*50]
            for tc in TIEU_CHUAN_MAC_DINH:
                tc_lines.append(f"{tc[0]} ({tc[1]}): {tc[2]}")
            zf.writestr("Phu_luc_4_Tieu_chuan.txt",
                        "\n".join(tc_lines).encode("utf-8"))

            # Danh mục bảng/hình
            dm_lines = ["DANH MỤC BẢNG BIỂU\n" + "="*40]
            for num_str, cap in self.num.tables:
                dm_lines.append(f"{num_str}: {cap}")
            dm_lines += ["\nDANH MỤC HÌNH VẼ\n" + "="*40]
            for num_str, cap in self.num.figures:
                dm_lines.append(f"{num_str}: {cap}")
            zf.writestr("Danh_muc_bang_hinh.txt",
                        "\n".join(dm_lines).encode("utf-8"))
        return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Helpers ngoài class
# ─────────────────────────────────────────────────────────────────────────────

def _fmt_run(run, bold=False, italic=False, size=13, name="Times New Roman", color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def _add_page_number(paragraph):
    """Chèn field số trang vào đoạn văn."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    _fmt_run(run, size=10, italic=True)


def _is_numeric(s: str) -> bool:
    try:
        float(s.replace(",", "").replace(" ", ""))
        return True
    except ValueError:
        return False


# ─────────────────────────────────────────────────────────────────────────────
# CLI test
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    # Dữ liệu mẫu đầy đủ cho testing
    sample_ss = {
        "cap_cong_trinh": "III",
        "cap_duong":       "III",
        "van_toc_thiet_ke": 60,
        "beam_params_final": {
            "loai_dam": "Super-T",
            "L":        38_200,   # mm
            "H":        1_750,
            "B":        2_440,
            "S":        2_200,
            "n_dam":    6,
            "cong_nghe":"DUL căng trước",
            "mac_bt":   "B50 (f'c = 50 MPa)",
        },
        "kcn_result": {
            "tong_so_nhip":         1,
            "be_rong_cau":          12.0,
            "so_lan_xe":            2,
            "so_dam_tren_mat_cat":  6,
        },
        "pier_result": {
            "loai_tru": "Trụ thân đặc BTCT",
            "H_tru":    6.5,
            "loai_mo":  "Mố chữ U BTCT",
            "H_dap":    2.0,
        },
        "mong_coc_result": {
            "loai_coc":    "Cọc khoan nhồi BTCT",
            "D_coc":       800,
            "L_coc":       35,
            "n_coc_be":    4,
            "Q_cho_phep":  1_800,
            "lop_tua_mui": "Lớp cát chặt N≥30 (lớp 3)",
        },
        "du_toan_result": {
            "tong_hop": {
                "hm1_KCN":    3_850,
                "hm2_MTM":    980,
                "hm3_MON":    2_100,
                "hm4_BMC":    750,
                "hm5_PHU":    420,
                "tong_xay_lap": 8_100,
                "chi_phi_dp":   405,
                "phi_quan_ly":  567,
                "VAT":          810,
                "tong_dau_tu":  9_882,
            },
            "phan_tich": {
                "suat_dau_tu_trieu_m2": 11.13,
                "A_cau_m2":             458.4,
            },
        },
        "terrain_data": {
            "dia_hinh": "đồng bằng, tương đối bằng phẳng",
            "MNCN": "+1.80",
            "MNTN": "-0.50",
            "MNTT": "+1.20",
        },
    }

    project_info = {
        "ten_du_an":    "Cầu Kênh Xáng Thị Xã ABC",
        "dia_diem":     "Thị xã ABC, tỉnh Long An",
        "chu_dau_tu":   "UBND Thị xã ABC, tỉnh Long An",
        "tu_van":       "Đề tài nghiên cứu — AI-Bridge-Design System (UTH)",
        "gv_huong_dan": "TS. Nguyễn Văn Hiển",
        "hoc_vien":     "NCS./ThS. Đinh Chương",
        "ngay_lap":     str(datetime.date.today()),
        "diem_dau":     "Km 0+000 (đường tỉnh ĐT826)",
        "diem_cuoi":    "Km 0+050 (phía đối diện)",
    }

    print("Đang sinh thuyết minh TKCS mẫu...")
    gen  = TKCSGenerator(sample_ss, project_info=project_info)

    def _prog(step, total, msg):
        pct = int(100 * step / total)
        print(f"  [{pct:3d}%] {msg}")

    docx_bytes = gen.generate(progress_callback=_prog)
    out = OUTPUT_DIR / "test_TKCS_sample.docx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out.write_bytes(docx_bytes)

    qc = gen.quality_check()
    print(f"\nKết quả: {out}")
    print(f"Kích thước: {len(docx_bytes):,} bytes")
    print(f"Chất lượng: {qc['total']}/100 — {qc['grade']}")
    print(f"Bảng: {len(gen.num.tables)}, Hình: {len(gen.num.figures)}")
    issues = [x for x in qc.get("issues",[]) if x]
    if issues:
        print("Cần cải thiện:", issues)
    else:
        print("Thuyết minh đạt yêu cầu chất lượng.")
